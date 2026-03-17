from datetime import datetime
from pathlib import Path
import json
import re

try:
    from docx import Document
    DOCX_LIBRARY_AVAILABLE = True
except ImportError:
    DOCX_LIBRARY_AVAILABLE = False
    print("AVISO: Biblioteca python-docx não encontrada. Instale com:")
    print("  pip install python-docx")


def tabela_para_markdown(tabela):
    """
    Converte uma tabela do python-docx para formato markdown.
    """
    if not tabela.rows:
        return ""
    
    linhas_md = []
    
    # Processa cada linha da tabela
    for i, row in enumerate(tabela.rows):
        celulas = []
        for cell in row.cells:
            texto_celula = cell.text.strip().replace('\n', ' ').replace('|', '\\|')
            # Se célula vazia, adiciona espaço para manter formato
            if not texto_celula:
                texto_celula = ' '
            celulas.append(texto_celula)
        
        if celulas:  # Só adiciona se houver células
            linha_md = '| ' + ' | '.join(celulas) + ' |'
            linhas_md.append(linha_md)
            
            # Adiciona linha de separação após o cabeçalho
            if i == 0:
                separador = '| ' + ' | '.join(['---'] * len(celulas)) + ' |'
                linhas_md.append(separador)
    
    return '\n'.join(linhas_md) if linhas_md else ""


def extrair_bloco_por_marcadores(docx_path, marcador_inicio, marcador_fim):
    """
    Extrai conteúdo de um DOCX entre marcadores de início e fim.
    Procura pelos marcadores no texto e extrai parágrafos e tabelas entre eles.
    """
    doc = Document(docx_path)
    
    # Extrai todos os elementos (parágrafos e tabelas) em ordem
    elementos = []
    
    # Mapeia elementos para seus índices nas listas do documento
    para_map = {para._element: i for i, para in enumerate(doc.paragraphs)}
    tabela_map = {tabela._element: i for i, tabela in enumerate(doc.tables)}
    
    # Percorre todos os elementos do documento em ordem usando XML
    for elemento in doc.element.body:
        tag = elemento.tag
        
        # Verifica se é um parágrafo (namespace w:p)
        if '}p' in tag or tag.endswith('p'):
            if elemento in para_map:
                para = doc.paragraphs[para_map[elemento]]
                texto = para.text.strip()
                if texto:
                    elementos.append(('paragrafo', texto))
        
        # Verifica se é uma tabela (namespace w:tbl)
        elif '}tbl' in tag or tag.endswith('tbl'):
            if elemento in tabela_map:
                tabela = doc.tables[tabela_map[elemento]]
                tabela_md = tabela_para_markdown(tabela)
                if tabela_md:
                    elementos.append(('tabela', tabela_md))
    
    if not elementos:
        return ""
    
    # Procura o marcador de início
    inicio_idx = None
    fim_idx = None
    
    # Normaliza o marcador de início (remove espaços extras, converte para maiúsculas)
    marcador_inicio_normalizado = ' '.join(marcador_inicio.strip().upper().split())
    
    # Procura marcador de início (busca parcial e case-insensitive)
    for idx, (tipo, conteudo) in enumerate(elementos):
        if tipo == 'paragrafo':
            texto_normalizado = ' '.join(conteudo.strip().upper().split())
            # Verifica se o marcador está contido no parágrafo
            # Usa busca parcial para ser mais flexível
            if marcador_inicio_normalizado in texto_normalizado:
                inicio_idx = idx
                break
    
    if inicio_idx is None:
        # Tenta busca mais flexível: procura por palavras-chave principais
        palavras_chave = marcador_inicio_normalizado.split()
        if len(palavras_chave) > 3:
            # Usa as primeiras palavras-chave significativas
            palavras_principais = palavras_chave[:4]
            marcador_flexivel = ' '.join(palavras_principais)
            
            for idx, (tipo, conteudo) in enumerate(elementos):
                if tipo == 'paragrafo':
                    texto_normalizado = ' '.join(conteudo.strip().upper().split())
                    if marcador_flexivel in texto_normalizado:
                        inicio_idx = idx
                        break
    
    if inicio_idx is None:
        return f"(Marcador de início não encontrado: {marcador_inicio})"
    
    # Procura marcador de fim (se fornecido)
    if marcador_fim and marcador_fim.strip():
        marcador_fim_normalizado = ' '.join(marcador_fim.strip().upper().split())
        
        for idx in range(inicio_idx + 1, len(elementos)):
            tipo, conteudo = elementos[idx]
            if tipo == 'paragrafo':
                texto_normalizado = ' '.join(conteudo.strip().upper().split())
                # Verifica se o marcador está contido no parágrafo
                if marcador_fim_normalizado in texto_normalizado:
                    fim_idx = idx
                    break
        
        # Se não encontrou com busca exata, tenta busca flexível
        if fim_idx is None:
            palavras_chave_fim = marcador_fim_normalizado.split()
            if len(palavras_chave_fim) > 2:
                palavras_principais_fim = palavras_chave_fim[:3]
                marcador_fim_flexivel = ' '.join(palavras_principais_fim)
                
                for idx in range(inicio_idx + 1, len(elementos)):
                    tipo, conteudo = elementos[idx]
                    if tipo == 'paragrafo':
                        texto_normalizado = ' '.join(conteudo.strip().upper().split())
                        if marcador_fim_flexivel in texto_normalizado:
                            fim_idx = idx
                            break
    
    # Se não encontrou marcador de fim, usa até o final do documento
    if fim_idx is None:
        fim_idx = len(elementos)
    
    # Extrai os elementos no intervalo (inclui o marcador de início, exclui o de fim)
    elementos_selecionados = elementos[inicio_idx:fim_idx]
    
    if not elementos_selecionados:
        return "(Nenhum conteúdo encontrado entre os marcadores)"
    
    # Converte para texto markdown
    texto_completo = []
    
    for tipo, conteudo in elementos_selecionados:
        if tipo == 'paragrafo':
            texto_completo.append(conteudo)
        elif tipo == 'tabela':
            texto_completo.append('\n' + conteudo + '\n')
    
    return '\n'.join(texto_completo)


def criar_markdown_bloco(versao, bloco, texto, data_criacao):
    """
    Cria o conteúdo markdown para um bloco.
    """
    conteudo = f"""# Bloco {bloco} - Versão {versao}

**Data de Criação:** {data_criacao}  
**Versão:** {versao}

---

{texto}
"""
    return conteudo


def processar_versoes():
    """
    Processa todas as versões e extrai as regras de cada bloco.
    """
    if not DOCX_LIBRARY_AVAILABLE:
        print("ERRO: Instale a biblioteca python-docx antes de continuar.")
        return

    base_dir = Path(__file__).parent
    versoes_json = base_dir / "versoes.json"
    output_dir = base_dir / "documentacao_blocos"
    output_dir.mkdir(exist_ok=True)

    print(f"Lendo {versoes_json}...")
    with open(versoes_json, 'r', encoding='utf-8') as f:
        dados = json.load(f)

    versoes = dados.get("versoes", [])
    print(f"Encontradas {len(versoes)} versões para processar.\n")

    for idx, versao_info in enumerate(versoes, 1):
        versao = versao_info["versao"]
        data_criacao = versao_info.get("data_criacao", datetime.now().strftime("%Y-%m-%d"))
        caminho_arquivo = versao_info["caminho_arquivo"]
        blocos = versao_info.get("blocos", [])
        
        print(f"[{idx}/{len(versoes)}] Processando versão {versao}...")
        
        docx_path = base_dir / caminho_arquivo
        
        if not docx_path.exists():
            print(f"  AVISO: Arquivo DOCX não encontrado: {docx_path}")
            continue
        
        versao_dir = output_dir / versao
        versao_dir.mkdir(exist_ok=True)
        
        if not blocos:
            print(f"  AVISO: Nenhum bloco definido para a versão {versao}")
            continue
        
        for bloco_info in blocos:
            bloco = bloco_info["bloco"]
            marcador_inicio = bloco_info.get("inicio", "")
            marcador_fim = bloco_info.get("fim", "")
            
            if not marcador_inicio:
                print(f"  AVISO: Bloco {bloco} não tem marcador de início definido")
                continue
            
            try:
                print(f"  Extraindo Bloco {bloco}...")
                print(f"    Início: {marcador_inicio[:60]}...")
                print(f"    Fim: {marcador_fim[:60] if marcador_fim else 'Final do documento'}...")
                
                texto = extrair_bloco_por_marcadores(docx_path, marcador_inicio, marcador_fim)
                
                if not texto.strip() or texto.startswith("("):
                    print(f"    AVISO: {texto}")
                    continue
                
                conteudo_md = criar_markdown_bloco(versao, bloco, texto, data_criacao)
                
                arquivo_md = versao_dir / f"bloco_{bloco}.md"
                with open(arquivo_md, 'w', encoding='utf-8') as f:
                    f.write(conteudo_md)
                
                num_caracteres = len(texto)
                num_linhas = texto.count('\n')
                print(f"    ✓ Salvo em: {arquivo_md.relative_to(base_dir)} (~{num_caracteres} caracteres, {num_linhas} linhas)")
            except Exception as e:
                print(f"    ERRO ao processar Bloco {bloco}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        print()
    
    print(f"\n✓ Processamento concluído! Documentação salva em: {output_dir.relative_to(base_dir)}")


if __name__ == "__main__":
    processar_versoes()
